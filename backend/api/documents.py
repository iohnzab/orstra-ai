import uuid
import io
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session

from database.connection import get_db
from database.models import Document, DocumentChunk, Agent
from api.auth import get_current_user, User
from utils.logger import get_logger

router = APIRouter(prefix="/documents", tags=["documents"])
logger = get_logger(__name__)


def _split_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks without any external dependencies."""
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap
    return chunks


def extract_text(filename: str, content: bytes) -> str:
    """Extract text from uploaded file."""
    if filename.lower().endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            logger.error("pdf_extract_error", error=str(e))
            return ""

    elif filename.lower().endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs)
        except Exception as e:
            logger.error("docx_extract_error", error=str(e))
            return ""

    else:
        try:
            return content.decode("utf-8", errors="replace")
        except Exception:
            return ""


@router.post("/upload")
async def upload_document(
    agent_id: str = Form(...),
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    agent = db.query(Agent).filter(Agent.id == agent_id, Agent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    content = await file.read()
    text = extract_text(file.filename or "file.txt", content)

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    # Split into chunks (simple built-in splitter — no external deps)
    chunks = _split_text(text, chunk_size=500, overlap=50)

    # Create document record
    doc = Document(
        id=uuid.uuid4(),
        agent_id=uuid.UUID(agent_id),
        filename=file.filename or "document.txt",
        content=text[:5000],  # Store first 5k chars
        chunk_count=len(chunks),
    )
    db.add(doc)
    db.flush()

    # Store chunks (no embeddings — using PostgreSQL FTS for search)
    for i, chunk_text in enumerate(chunks):
        chunk = DocumentChunk(
            id=uuid.uuid4(),
            document_id=doc.id,
            agent_id=uuid.UUID(agent_id),
            content=chunk_text,
            embedding=None,
            chunk_index=i,
        )
        db.add(chunk)

    db.commit()

    return {
        "id": str(doc.id),
        "filename": doc.filename,
        "chunk_count": len(chunks),
        "agent_id": agent_id,
    }


@router.get("/{agent_id}")
def list_documents(
    agent_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    agent = db.query(Agent).filter(Agent.id == agent_id, Agent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=404, detail="Agent not found")

    docs = db.query(Document).filter(Document.agent_id == agent_id).all()
    return [
        {
            "id": str(d.id),
            "filename": d.filename,
            "chunk_count": d.chunk_count,
            "created_at": d.created_at.isoformat(),
        }
        for d in docs
    ]


@router.delete("/{doc_id}", status_code=204)
def delete_document(
    doc_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Verify ownership
    agent = db.query(Agent).filter(Agent.id == doc.agent_id, Agent.user_id == current_user.id).first()
    if not agent:
        raise HTTPException(status_code=403, detail="Access denied")

    db.delete(doc)
    db.commit()
